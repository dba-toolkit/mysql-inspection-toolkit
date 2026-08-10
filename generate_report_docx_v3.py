#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate a professional MySQL inspection DOCX from report_model.json.

The generator is intentionally presentation-only: health scoring, risk
classification, metric validity and missing-value decisions belong to the
analyzer. This script renders that stable contract without recalculation.
"""
from __future__ import annotations

import argparse
import json
import math
import re
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


GENERATOR_VERSION = "3.1.0"
FONT_EN = "Calibri"
FONT_CN = "Microsoft YaHei"
COLOR_NAVY = "000000"
COLOR_BLUE = "000000"
COLOR_DARK_BLUE = "000000"
COLOR_TEXT = "111111"
COLOR_MUTED = "666666"
COLOR_BORDER = "B7B7B7"
COLOR_HEADER = "F2F2F2"
COLOR_ALT = "FAFAFA"
COLOR_CALLOUT = "F7F7F7"
COLOR_HIGH = "9B1C1C"
COLOR_MEDIUM = "B26A00"
COLOR_LOW = "6B5A00"
COLOR_GOOD = "222222"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120

SEVERITY_CN = {"critical": "严重", "high": "高", "medium": "中", "low": "低", "info": "提示"}
PRIORITY_LABELS = {"P1": "立即处理", "P2": "计划处理", "P3": "持续优化"}


def load_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as stream:
        value = json.load(stream)
    if not isinstance(value, dict):
        raise ValueError(f"JSON root must be an object: {path}")
    if value.get("generator_contract") != "mysql_inspection_report_model":
        raise ValueError("input is not a mysql_inspection report_model.json")
    return value


def text(value: Any, missing: str = "未采集") -> str:
    if value is None:
        return missing
    if value == "":
        return "-"
    return str(value)


def number(value: Any, digits: int = 1, suffix: str = "", missing: str = "未采集") -> str:
    if value is None or value == "":
        return missing
    try:
        parsed = float(value)
    except (TypeError, ValueError):
        return str(value)
    if not math.isfinite(parsed):
        return missing
    rendered = f"{parsed:,.{digits}f}"
    if digits:
        rendered = rendered.rstrip("0").rstrip(".")
    return rendered + suffix


def bytes_text(value: Any, missing: str = "未采集") -> str:
    if value is None or value == "":
        return missing
    try:
        amount = float(value)
    except (TypeError, ValueError):
        return str(value)
    if amount < 0 or not math.isfinite(amount):
        return missing
    units = ["B", "KB", "MB", "GB", "TB", "PB"]
    index = 0
    while amount >= 1024 and index < len(units) - 1:
        amount /= 1024
        index += 1
    return f"{amount:,.2f} {units[index]}"


def datetime_text(value: Any) -> str:
    raw = text(value)
    try:
        return datetime.fromisoformat(raw).strftime("%Y-%m-%d %H:%M:%S %z")
    except ValueError:
        return raw


def safe_filename_part(value: Any) -> str:
    """Return a Windows-safe filename component."""
    rendered = str(value or "").strip()
    rendered = re.sub(r'[<>:"/\\|?*\x00-\x1f]', "_", rendered)
    rendered = re.sub(r"\s+", " ", rendered).strip(" .")
    return rendered


def default_output_path(model_path: Path, model: dict[str, Any], target: str = "") -> Path:
    """Build: IP-hostname 系统名.docx or IP-hostname MySQL数据库巡检报告.docx."""
    overview = model.get("overview") or {}
    cover = model.get("cover") or {}
    ip_address = safe_filename_part(overview.get("ip"))
    hostname = safe_filename_part(
        overview.get("host") or cover.get("inspection_target")
    )
    identity = "_".join(part for part in (ip_address, hostname) if part)
    if not identity:
        identity = "MySQL"
    suffix = target.strip() if target.strip() else "MySQL数据库巡检报告"
    return model_path.parent / f"{identity}_{suffix}.docx"


def status_text(value: Any) -> str:
    mapping = {
        "ok": "通过", "complete": "完整", "usable": "可用", "unusable": "不可用",
        "success": "成功", "failed": "失败", "partial": "部分可用",
    }
    raw = text(value)
    return mapping.get(raw.lower(), raw)


def clean_limitation(value: Any) -> str:
    raw = str(value)
    if "system.filesystems" in raw:
        return "文件系统主体信息已取得，但 /run/user/0/gvfs 挂载点不可访问"
    if "system.inodes" in raw:
        return "inode 主体信息已取得，但 /run/user/0/gvfs 挂载点不可访问"
    if "system.sar_history" in raw:
        return "SAR 原始数据已导出，但历史覆盖不足，由分析器重新校验覆盖范围"
    return raw


def set_run_font(run, size: float = 11, bold: bool = False, color: str = COLOR_TEXT,
                 italic: bool = False, font_en: str = FONT_EN, font_cn: str = FONT_CN) -> None:
    run.font.name = font_en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), font_en)
    rpr.rFonts.set(qn("w:hAnsi"), font_en)
    rpr.rFonts.set(qn("w:eastAsia"), font_cn)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color: str = COLOR_BORDER, size: int = 5) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = TABLE_INDENT_DXA) -> None:
    if sum(widths_dxa) != CONTENT_DXA:
        raise ValueError(f"table widths must total {CONTENT_DXA} DXA: {widths_dxa}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    cached = OxmlElement("w:t")
    cached.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, cached, end])
    set_run_font(run, size=9, color=COLOR_MUTED)


class ReportGenerator:
    def __init__(self, model_path: Path, output_path: Path, customer: str = "",
                 author: str = "", reviewer: str = "", report_version: str = "",
                 logo_path: Path | None = None, target: str = "", company: str = "") -> None:
        self.model_path = model_path.resolve()
        self.root = self.model_path.parent
        self.output_path = output_path.resolve()
        self.model = load_json(self.model_path)
        self.cover = self.model.get("cover") or {}
        self.control = self.model.get("document_control") or {}
        self.customer = customer or text(self.control.get("customer"), "待填写")
        self.company = company or ""
        self.target = target or text(self.cover.get("inspection_target"), "待填写")
        self.author = author or "自动生成"
        self.reviewer = reviewer or "待填写"
        self.report_version = report_version or text(self.control.get("report_version"), "V1.0")
        self.logo_path = logo_path
        self.doc = Document()
        self._configure_document()
        self._setup_first_page_header()

    def _setup_first_page_header(self) -> None:
        """Company name in first-page header only; no header on subsequent pages."""
        section = self.doc.sections[0]
        # Clear default header (pages 2+)
        default = section.header
        default.is_linked_to_previous = False
        for p in default.paragraphs:
            p.text = ""
        # First page header — company name
        header = section.first_page_header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(self.company)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string(COLOR_MUTED)

    def _configure_document(self) -> None:
        section = self.doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)
        section.different_first_page_header_footer = True

        normal = self.doc.styles["Normal"]
        normal.font.name = FONT_EN
        normal.font.size = Pt(11)
        normal.font.color.rgb = RGBColor.from_string(COLOR_TEXT)
        rpr = normal._element.get_or_add_rPr()
        rpr.rFonts.set(qn("w:ascii"), FONT_EN)
        rpr.rFonts.set(qn("w:hAnsi"), FONT_EN)
        rpr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.1

        heading_tokens = {
            1: (16, COLOR_BLUE, 16, 8),
            2: (13, COLOR_BLUE, 12, 6),
            3: (12, COLOR_DARK_BLUE, 8, 4),
        }
        for level, (size, color, before, after) in heading_tokens.items():
            style = self.doc.styles[f"Heading {level}"]
            style.font.name = FONT_EN
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            srpr = style._element.get_or_add_rPr()
            srpr.rFonts.set(qn("w:ascii"), FONT_EN)
            srpr.rFonts.set(qn("w:hAnsi"), FONT_EN)
            srpr.rFonts.set(qn("w:eastAsia"), FONT_CN)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        set_run_font(header.add_run("MySQL 数据库巡检分析报告"), size=9, color=COLOR_MUTED)
        footer = section.footer.paragraphs[0]
        add_page_field(footer)
        first_header = section.first_page_header.paragraphs[0]
        first_header.text = ""
        first_footer = section.first_page_footer.paragraphs[0]
        first_footer.text = ""

        properties = self.doc.core_properties
        properties.title = text(self.cover.get("title"), "MySQL数据库巡检分析报告")
        properties.subject = "数据库巡检"
        properties.author = self.author
        properties.comments = ""

    def paragraph(self, value: Any = "", *, size: float = 11, bold: bool = False,
                  color: str = COLOR_TEXT, italic: bool = False,
                  align: WD_ALIGN_PARAGRAPH | None = None, before: float = 0,
                  after: float = 6, keep_next: bool = False) -> Any:
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.line_spacing = 1.1
        p.paragraph_format.keep_with_next = keep_next
        set_run_font(p.add_run(text(value, "")), size=size, bold=bold, color=color, italic=italic)
        return p

    def heading(self, value: str, level: int = 1) -> Any:
        return self.doc.add_heading(value, level=level)

    def page_break(self) -> None:
        self.doc.add_page_break()

    def horizontal_rule(self, *, inset: float = 0, color: str = "000000",
                        size: int = 8, before: float = 0, after: float = 0) -> Any:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(inset)
        p.paragraph_format.right_indent = Pt(inset)
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(size))
        bottom.set(qn("w:color"), color)
        bottom.set(qn("w:space"), "1")
        borders.append(bottom)
        p_pr.append(borders)
        return p

    def note_box(self, label: str, value: str, fill: str = COLOR_CALLOUT, accent: str = COLOR_BLUE) -> None:
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(8)
        p.paragraph_format.right_indent = Pt(8)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        p_pr.append(shd)
        borders = OxmlElement("w:pBdr")
        for edge in ("top", "left", "bottom", "right"):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), "single")
            node.set(qn("w:sz"), "4")
            node.set(qn("w:color"), COLOR_BORDER)
            node.set(qn("w:space"), "6")
            borders.append(node)
        p_pr.append(borders)
        set_run_font(p.add_run(f"{label}  "), size=10.5, bold=True, color=accent)
        set_run_font(p.add_run(value), size=10.5, color=COLOR_TEXT)

    def table(self, headers: list[str], rows: Iterable[Iterable[Any]], widths_dxa: list[int],
              *, compact: bool = False, severity_column: int | None = None) -> Any:
        table = self.doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        header = table.rows[0]
        set_repeat_header(header)
        set_row_cant_split(header)
        for index, value in enumerate(headers):
            cell = header.cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_shading(cell, COLOR_HEADER)
            set_cell_border(cell)
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            set_run_font(p.add_run(text(value, "")), size=9 if compact else 9.5, bold=True, color=COLOR_NAVY)
        for row_index, values in enumerate(rows):
            row = table.add_row()
            set_row_cant_split(row)
            for column, value in enumerate(list(values)):
                cell = row.cells[column]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell)
                set_cell_margins(cell, 90, 120, 90, 120)
                if row_index % 2 == 1:
                    set_cell_shading(cell, COLOR_ALT)
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                if column == 0 and len(headers) > 2:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rendered = text(value)
                color = COLOR_TEXT
                bold = False
                if severity_column == column:
                    lowered = rendered.lower()
                    if rendered in {"高", "严重"} or lowered in {"high", "critical"}:
                        color, bold = COLOR_HIGH, True
                    elif rendered == "中" or lowered == "medium":
                        color, bold = COLOR_MEDIUM, True
                    elif rendered == "低" or lowered == "low":
                        color, bold = COLOR_LOW, True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p.add_run(rendered), size=8.7 if compact else 9.5, bold=bold, color=color)
        set_table_geometry(table, widths_dxa)
        self.paragraph("", after=2)
        return table

    @staticmethod
    def _dynamic_widths(headers: list[str], rows: list[dict[str, Any]]) -> list[int]:
        if not headers:
            return []
        minimum = 650 if len(headers) >= 7 else 800 if len(headers) >= 5 else 1000
        lengths: list[float] = []
        for header in headers:
            values = [str(row.get(header, "") or "") for row in rows[:30]]
            longest = max([len(str(header)), *(min(len(value), 36) for value in values)])
            numeric = all(
                not value or value.replace(".", "", 1).replace("%", "", 1).replace("-", "", 1).isdigit()
                for value in values
            )
            lengths.append(max(4.0, longest * (0.68 if numeric else 1.0)))
        remaining = CONTENT_DXA - minimum * len(headers)
        if remaining < 0:
            base = CONTENT_DXA // len(headers)
            widths = [base] * len(headers)
            widths[-1] += CONTENT_DXA - sum(widths)
            return widths
        total_weight = sum(lengths) or len(headers)
        widths = [minimum + int(remaining * weight / total_weight) for weight in lengths]
        widths[-1] += CONTENT_DXA - sum(widths)
        return widths

    def data_table(self, rows: list[dict[str, Any]], *, note: str = "") -> Any | None:
        if not rows:
            return None
        headers: list[str] = []
        for row in rows:
            for key in row:
                if key not in headers:
                    headers.append(str(key))
        widths = self._dynamic_widths(headers, rows)
        compact = len(headers) >= 5
        rendered_rows = [[row.get(header) for header in headers] for row in rows]
        table = self.table(headers, rendered_rows, widths, compact=compact)
        if note:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(7)
            set_run_font(p.add_run(f"说明：{note}"), size=8.5, color=COLOR_MUTED, italic=True)
        return table

    def analysis_block(self, analysis: dict[str, Any]) -> None:
        status = str(analysis.get("status") or "not_evaluated")
        status_cn = {
            "normal": "正常",
            "attention": "关注",
            "risk": "风险",
            "not_evaluated": "证据不足",
            "not_applicable": "不适用",
        }.get(status, status)
        color = COLOR_HIGH if status == "risk" else COLOR_MEDIUM if status == "attention" else COLOR_TEXT
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Pt(4)
        set_run_font(p.add_run(f"分析结论（{status_cn}）："), size=10.5, bold=True, color=color)
        set_run_font(p.add_run(text(analysis.get("conclusion"))), size=10.5, color=COLOR_TEXT)
        evidence = analysis.get("evidence") or []
        if evidence:
            ep = self.doc.add_paragraph()
            ep.paragraph_format.left_indent = Pt(4)
            ep.paragraph_format.space_before = Pt(0)
            ep.paragraph_format.space_after = Pt(3)
            set_run_font(ep.add_run("判断依据："), size=9.5, bold=True, color=COLOR_TEXT)
            set_run_font(ep.add_run("；".join(str(item) for item in evidence)), size=9.5, color=COLOR_MUTED)
        recommendation = str(analysis.get("recommendation") or "").strip()
        if recommendation:
            rp = self.doc.add_paragraph()
            rp.paragraph_format.left_indent = Pt(4)
            rp.paragraph_format.space_before = Pt(0)
            rp.paragraph_format.space_after = Pt(8)
            set_run_font(rp.add_run("建议："), size=10.5, bold=True, color=COLOR_TEXT)
            set_run_font(rp.add_run(recommendation), size=10.5, color=COLOR_TEXT)

    def render_inspection_item(self, item: dict[str, Any], heading_number: str) -> None:
        self.heading(f"{heading_number} {text(item.get('title'))}", 2)
        source = str(item.get("source") or "").strip()
        collection = item.get("collection") or {}
        status = status_text(collection.get("status"))
        source_line = f"数据来源：{source or '报告模型'}；采集状态：{status}"
        if collection.get("row_count") is not None:
            source_line += f"；原始记录：{collection.get('row_count')} 条"
        reason = str(collection.get("reason") or "").strip()
        if reason:
            source_line += f"；说明：{reason}"
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        set_run_font(p.add_run(source_line), size=8.5, color=COLOR_MUTED)
        display = item.get("display") or {}
        rows = display.get("rows") or []
        if rows:
            self.data_table(rows, note=str(display.get("note") or ""))
        else:
            empty_text = "本次采集无相关记录。"
            if collection.get("status") not in {"empty", "not_applicable"}:
                empty_text = "本项没有可展示的结构化记录。"
            self.note_box("数据状态", empty_text, fill=COLOR_CALLOUT, accent=COLOR_TEXT)
        self.analysis_block(item.get("analysis") or {})

    def picture(self, chart: dict[str, Any], caption: str, width_inches: float = 6.15) -> bool:
        if chart.get("status") != "generated" or not chart.get("file"):
            self.note_box("图表未生成", f"{caption}：{text(chart.get('reason'), '数据点不足')}。")
            return False
        path = self.root / str(chart["file"])
        if not path.exists():
            self.note_box("图表缺失", f"{caption}：未找到 {path.name}。")
            return False
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        shape = p.add_run().add_picture(str(path), width=Inches(width_inches))
        shape._inline.docPr.set("descr", caption)
        shape._inline.docPr.set("title", caption)
        cp = self.doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after = Pt(8)
        set_run_font(cp.add_run(caption), size=9, color=COLOR_MUTED)
        return True

    def build_cover(self) -> None:
        # Top thick frame line
        self.horizontal_rule(inset=0, color="000000", size=12, before=38, after=72)

        # Inspection target — the main object being inspected
        self.paragraph(self.target, size=28, bold=True, color="111111",
                       align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        # Thin divider
        self.horizontal_rule(inset=150, color=COLOR_MUTED, size=5, after=8)

        # Report type
        self.paragraph(text(self.cover.get("title"), "MySQL数据库巡检分析报告"), size=20,
                       color=COLOR_NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
        # Thin divider
        self.horizontal_rule(inset=150, color=COLOR_MUTED, size=5, after=36)

        # Logo
        if self.logo_path and self.logo_path.exists():
            logo_p = self.doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_p.paragraph_format.space_before = Pt(8)
            logo_p.paragraph_format.space_after = Pt(36)
            try:
                run = logo_p.add_run()
                run.add_picture(str(self.logo_path), width=Inches(2.5))
            except Exception:
                pass

        # Bottom thick frame line (above info table)
        self.horizontal_rule(inset=0, color="000000", size=12, before=12, after=12)

        # Cover info table
        self.table(
            ["项目", "内容"],
            [
                ["巡检对象", self.target],
                ["数据库版本", "MySQL " + text(self.cover.get("database_version"))],
                ["报告版本", self.report_version],
                ["巡检日期", text(self.cover.get("inspection_date"))],
            ],
            [2500, 6860],
        )
        self.page_break()

    def build_document_control(self) -> None:
        self.heading("文档控制", 1)
        self.table(
            ["项目", "内容"],
            [
                ["客户", self.customer],
                ["数据库", text(self.control.get("database"), "MySQL")],
                ["报告版本", self.report_version],
                ["生成时间", datetime_text(self.control.get("generated_at"))],
                ["编制", self.author],
                ["复核", self.reviewer],
            ],
            [2500, 6860],
        )
        self.heading("修订记录", 2)
        self.table(
            ["版本", "日期", "修订说明", "编制"],
            [[self.report_version, text(self.cover.get("inspection_date")), "首次自动生成", self.author]],
            [1200, 1900, 4160, 2100],
        )
        self.note_box("保密说明", "本报告仅供客户授权人员用于数据库运行状况评估和整改决策。")
        self.page_break()
        self.heading("目录", 1)
        contents = [
            "1. 巡检总结", "2. 巡检概述", "3. 系统与环境检查", "4. 操作系统性能检查",
            "5. MySQL 实例与组件检查", "6. MySQL 参数配置检查", "7. MySQL 运行状态检查",
            "8. 容量与对象检查", "9. SQL、等待与文件 I/O 检查", "10. 账号与权限检查",
            "11. 日志、备份与复制检查", "12. 风险清单与整改建议", "13. 附录",
        ]
        for item in contents:
            self.paragraph(item, size=10.5, color=COLOR_TEXT, after=3)
        self.page_break()

    def _chart_map(self, section: dict[str, Any]) -> dict[str, dict[str, Any]]:
        return {str(chart.get("chart_id")): chart for chart in section.get("charts") or []}

    def build_summary_v31(self) -> None:
        health = self.model.get("health_assessment") or {}
        counts = health.get("counts") or {}
        score = health.get("score")
        conclusions = self.model.get("comprehensive_conclusions") or []
        self.heading("1. 巡检总结", 1)
        self.table(
            ["总体健康", "高风险", "中风险", "低风险"],
            [[
                f"{number(score, 0)} / 100",
                number(counts.get("high"), 0),
                number(counts.get("medium"), 0),
                number(counts.get("low"), 0),
            ]],
            [3000, 2120, 2120, 2120],
        )
        if conclusions:
            rows = []
            status_map = {
                "normal": "正常", "attention": "关注", "risk": "风险",
                "not_evaluated": "证据不足", "not_applicable": "不适用",
            }
            for index, item in enumerate(conclusions, 1):
                rows.append([
                    index,
                    status_map.get(str(item.get("status")), text(item.get("status"))),
                    f"{text(item.get('topic'))}：{text(item.get('conclusion'))}",
                ])
            self.heading("1.1 综合结论", 2)
            table = self.table(["序号", "状态", "结论内容"], rows, [800, 1200, 7360], compact=False)
            for row, item in zip(table.rows[1:], conclusions):
                status = str(item.get("status"))
                color = COLOR_HIGH if status == "risk" else COLOR_MEDIUM if status == "attention" else COLOR_TEXT
                for run in row.cells[1].paragraphs[0].runs:
                    set_run_font(run, size=9.5, bold=True, color=color)
        self.note_box(
            "结论口径",
            "健康评分、风险等级和检查结论均直接读取分析结果；Word 生成器仅负责呈现，不重新计算或改变判断。",
            fill=COLOR_CALLOUT,
            accent=COLOR_TEXT,
        )

    def build_overview_v31(self) -> None:
        overview = self.model.get("overview") or {}
        topology = self.model.get("topology") or {}
        self.page_break()
        self.heading("2. 巡检概述", 1)
        self.heading("2.1 巡检对象", 2)
        self.table(
            ["项目", "内容"],
            [
                ["主机", text(overview.get("host"))],
                ["IP 地址", text(overview.get("ip"))],
                ["数据库版本", text(overview.get("database_version"))],
                ["采集时间", datetime_text(overview.get("collection_time"))],
            ],
            [2500, 6860],
        )
        self.heading("2.2 架构拓扑", 2)
        nodes = topology.get("nodes") or []
        edges = topology.get("edges") or []
        if topology.get("mode") == "single_instance" or (len(nodes) == 1 and not edges):
            node = nodes[0] if nodes else {}
            self.table(
                ["拓扑模式", "节点", "地址", "角色"],
                [[
                    "单实例",
                    text(node.get("hostname") or node.get("instance_tag")),
                    f"{text(node.get('ip'))}:{text(node.get('port'))}",
                    text(node.get("role_observed")),
                ]],
                [1500, 2300, 2500, 3060],
            )
        else:
            self.table(
                ["节点", "地址", "角色", "版本"],
                [[
                    text(node.get("hostname") or node.get("instance_tag")),
                    f"{text(node.get('ip'))}:{text(node.get('port'))}",
                    text(node.get("role_observed")),
                    text(node.get("version")),
                ] for node in nodes],
                [2300, 2200, 2500, 2360],
                compact=True,
            )
            if edges:
                self.table(
                    ["源节点", "目标节点", "关系"],
                    [[
                        text(edge.get("source") or edge.get("source_node_id")),
                        text(edge.get("target") or edge.get("target_node_id")),
                        text(edge.get("type"), "复制"),
                    ] for edge in edges],
                    [3300, 3300, 2760],
                    compact=True,
                )

    def render_section_charts(self, section_id: str, main_number: int, subsection_number: int) -> None:
        if section_id == "system_performance":
            charts = self._chart_map(self.model.get("system_analysis") or {})
            definitions = [
                ("SYSTEM_CPU", "CPU 使用率趋势"),
                ("SYSTEM_MEMORY", "内存使用率趋势"),
                ("SYSTEM_DISK", "磁盘 I/O 趋势"),
                ("SYSTEM_NETWORK_REALTIME", "网络吞吐趋势"),
            ]
        elif section_id == "mysql_runtime":
            charts = self._chart_map(self.model.get("mysql_performance") or {})
            definitions = [
                ("MYSQL_QPS_TPS", "QPS 与 TPS 趋势"),
                ("MYSQL_THREADS", "连接与运行线程趋势"),
            ]
        else:
            return
        generated = [(chart_id, caption) for chart_id, caption in definitions if charts.get(chart_id, {}).get("status") == "generated"]
        if not generated:
            return
        self.heading(f"{main_number}.{subsection_number} 趋势图", 2)
        for index, (chart_id, caption) in enumerate(generated, 1):
            self.picture(charts.get(chart_id, {}), f"图 {main_number}-{index} {caption}", width_inches=5.85)

    def build_detailed_sections_v31(self) -> None:
        sections = self.model.get("inspection_sections") or []
        order = {
            "system_environment": 0,
            "system_performance": 1,
            "mysql_instance": 2,
            "mysql_configuration": 3,
            "mysql_runtime": 4,
            "capacity_objects": 5,
            "sql_io": 6,
            "security": 7,
            "logs_backup_replication": 8,
        }
        sections = sorted(sections, key=lambda section: order.get(str(section.get("section_id")), 99))
        main_number = 3
        for section in sections:
            self.page_break()
            self.heading(f"{main_number}. {text(section.get('title'))}", 1)
            items = section.get("items") or []
            if not items:
                self.note_box("数据状态", "本章节没有检查项。", accent=COLOR_TEXT)
            for item_index, item in enumerate(items, 1):
                self.render_inspection_item(item, f"{main_number}.{item_index}")
            self.render_section_charts(str(section.get("section_id")), main_number, len(items) + 1)
            main_number += 1

    def build_risk_and_plan_v31(self) -> None:
        findings = self.model.get("risk_register") or []
        plan = self.model.get("optimization_plan") or {}
        self.page_break()
        self.heading("12. 风险清单与整改建议", 1)
        self.heading("12.1 风险清单", 2)
        if findings:
            rows = []
            for finding in findings:
                evidence = "；".join(str(item) for item in finding.get("facts") or [])
                rows.append([
                    text(finding.get("finding_id")),
                    SEVERITY_CN.get(finding.get("severity"), text(finding.get("severity"))),
                    text(finding.get("title")),
                    evidence or text(finding.get("summary")),
                    text(finding.get("recommendation")),
                ])
            self.table(
                ["编号", "级别", "问题", "证据", "建议"],
                rows,
                [750, 750, 1900, 2700, 3260],
                compact=True,
                severity_column=1,
            )
        else:
            self.note_box("风险结论", "本次已评价规则未发现风险。", accent=COLOR_TEXT)
        self.heading("12.2 整改优先级", 2)
        for index, priority in enumerate(("P1", "P2", "P3"), 1):
            self.heading(f"12.2.{index} {priority} {PRIORITY_LABELS[priority]}", 3)
            items = plan.get(priority) or []
            if not items:
                self.paragraph("本优先级暂无整改项。", color=COLOR_MUTED)
                continue
            self.table(
                ["编号", "整改项", "建议措施"],
                [[
                    text(item.get("finding_id")),
                    text(item.get("title")),
                    text(item.get("recommendation")),
                ] for item in items],
                [1000, 2900, 5460],
                compact=True,
            )

    def build_appendix_v31(self) -> None:
        appendix = self.model.get("appendix") or {}
        window = appendix.get("collection_window") or {}
        history = window.get("history") or {}
        quality = appendix.get("data_quality") or {}
        evaluations = appendix.get("rule_evaluations") or []
        gaps = self.model.get("collection_gaps") or []
        self.page_break()
        self.heading("13. 附录", 1)
        self.heading("13.1 采集窗口与数据完整性", 2)
        self.table(
            ["项目", "内容"],
            [
                ["实时采集窗口", number(window.get("realtime_window_seconds"), 1, " 秒")],
                ["MySQL 采样点", number(window.get("mysql_sample_points"), 0)],
                ["短窗口", "是" if window.get("short_window") else "否"],
                ["SAR 历史状态", status_text(history.get("status"))],
                ["SAR 覆盖", f"{number(history.get('coverage_hours'), 2)} / {number(history.get('requested_hours'), 0)} 小时"],
                ["数据完整度", number(quality.get("score"), 1, "%")],
                ["完整度等级", text(quality.get("grade"))],
                ["采集包完整性", status_text((quality.get("integrity") or {}).get("status"))],
            ],
            [3000, 6360],
        )
        limitations = quality.get("limitations") or []
        if limitations:
            self.note_box(
                "数据限制",
                "；".join(clean_limitation(item) for item in limitations),
                fill="FFF8E8",
                accent=COLOR_MEDIUM,
            )
        if gaps:
            self.heading("13.2 数据缺口", 2)
            gap_status_cn = {
                "partial": "部分可用",
                "insufficient_history": "历史不足",
                "external_evidence_required": "需要外部证据",
                "permission_denied": "权限不足",
                "timeout": "超时",
                "error": "失败",
            }
            self.data_table([
                {
                    "采集项": gap.get("item_id"),
                    "状态": gap_status_cn.get(str(gap.get("status")), gap.get("status")),
                    "原因": gap.get("reason"),
                    "补充建议": gap.get("recommended_action"),
                }
                for gap in gaps
            ])
        self.heading("13.3 规则执行状态", 2)
        status_cn = {"triggered": "触发", "passed": "通过", "not_evaluated": "未判断", "not_applicable": "不适用"}
        counts: dict[str, int] = {}
        for item in evaluations:
            rule_status = str(item.get("status"))
            counts[rule_status] = counts.get(rule_status, 0) + 1
        self.table(
            ["状态", "数量", "含义"],
            [[status_cn[key], counts.get(key, 0), meaning] for key, meaning in (
                ("triggered", "形成风险项"),
                ("passed", "现有证据未触发阈值"),
                ("not_evaluated", "证据不足，不作通过结论"),
                ("not_applicable", "当前架构或环境不适用"),
            )],
            [2200, 1800, 5360],
        )
        self.heading("13.4 免责声明", 2)
        self.note_box(
            "免责声明",
            text(appendix.get("disclaimer")),
            fill="FFF8E8",
            accent=COLOR_MEDIUM,
        )

    def build(self) -> Path:
        self.build_cover()
        self.build_document_control()
        self.build_summary_v31()
        self.build_overview_v31()
        self.build_detailed_sections_v31()
        self.build_risk_and_plan_v31()
        self.build_appendix_v31()
        self.output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(self.output_path)
        return self.output_path


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a professional MySQL inspection DOCX from report_model.json.")
    parser.add_argument("input", help="report_model.json or its containing directory")
    parser.add_argument("--output", help="output .docx path")
    parser.add_argument("--customer", default="", help="customer name (e.g. 昆山市第一人民医院)")
    parser.add_argument("--company", default="", help="company name shown in page header")
    parser.add_argument("--target", default="", help="inspection target name override (e.g.  HIS数据库, 生产订单系统)")
    parser.add_argument("--author", default="", help="author/preparer")
    parser.add_argument("--reviewer", default="", help="reviewer")
    parser.add_argument("--report-version", default="", help="report version override")
    parser.add_argument("--logo", default="logo.png", help="path to logo image, PNG recommended (default: logo.png if exists)")
    return parser.parse_args()


def main() -> int:
    args = parse_args()
    source = Path(args.input).expanduser().resolve()
    model_path = source / "report_model.json" if source.is_dir() else source
    if not model_path.exists():
        print(f"ERROR: report_model.json not found: {model_path}", file=sys.stderr)
        return 1
    if args.output:
        output = Path(args.output).expanduser().resolve()
    else:
        model = load_json(model_path)
        output = default_output_path(model_path, model, args.target)
    try:
        logo_path = Path(args.logo).expanduser().resolve() if args.logo else None
        # silently skip default logo if file not found
        if logo_path is not None and not logo_path.exists():
            logo_path = None
        generated = ReportGenerator(
            model_path, output, customer=args.customer, author=args.author,
            reviewer=args.reviewer, report_version=args.report_version,
            logo_path=logo_path, target=args.target, company=args.company,
        ).build()
        print(json.dumps({"status": "success", "output": str(generated), "generator_version": GENERATOR_VERSION}, ensure_ascii=False))
        return 0
    except Exception as exc:
        print(f"ERROR: {type(exc).__name__}: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
